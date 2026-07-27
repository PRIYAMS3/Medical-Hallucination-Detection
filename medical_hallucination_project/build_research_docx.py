from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parent
OUT = PROJECT_DIR / "outputs" / "medical_hallucination_research_proposal.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "F4F6F9")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].add_run(value)

    doc.add_paragraph()


def build_doc() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(
        "Context-Aware Explainable Medical Hallucination Detection Using Retrieval-Augmented NLI, Fine-Grained Type Classification, and Clinical Severity Scoring"
    )

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Research proposal draft | Prepared from current project notes")

    doc.add_heading("1. Background", level=1)
    doc.add_paragraph(
        "Large language models are increasingly used for medical question answering, clinical summarization, biomedical literature assistance, and patient-facing health information. However, these models can produce hallucinations: fluent statements that are unsupported, contradicted by evidence, or clinically unsafe. In medicine, hallucinations are more serious than ordinary factual errors because they may affect diagnosis, treatment, medication advice, contraindications, patient safety, or clinical decision-making."
    )
    doc.add_paragraph(
        "Recent literature shows rapid progress but also fragmentation. MedHallu introduced a medical hallucination benchmark derived from PubMedQA and showed that strong models still struggle on hard medical hallucinations. Scientific hallucination work such as SciHal25 and From RAG to Reality highlights the usefulness of natural language inference. RAG hallucination studies such as RAGTruth, RAGChecker, HALT-RAG, ReDeEP, LettuceDetect, and RT4CHART explore retrieval grounding, span detection, interpretability, and claim-level verification."
    )

    doc.add_heading("2. Problem Statement", level=1)
    doc.add_paragraph(
        "Existing medical hallucination detection methods often focus on binary hallucination labels or benchmark-level evaluation without explaining the specific claim-level error, retrieving supporting or contradicting biomedical evidence, identifying the hallucination type, or estimating clinical risk. This creates a gap between model evaluation and real clinical usefulness."
    )

    doc.add_heading("3. Research Gap", level=1)
    doc.add_paragraph(
        "Recent work has advanced medical hallucination benchmarks, scientific NLI-based hallucination detection, RAG hallucination diagnostics, and clinical safety evaluation separately. However, no existing work provides a unified medical text hallucination detection framework that performs all of the following:"
    )
    add_numbered(
        doc,
        [
            "Decomposes medical answers into atomic claims.",
            "Retrieves biomedical evidence for each claim.",
            "Verifies each claim using NLI.",
            "Classifies hallucination type.",
            "Provides evidence-grounded explanations.",
            "Scores clinical severity based on risk.",
        ],
    )

    doc.add_heading("4. Aim and Objectives", level=1)
    doc.add_paragraph(
        "The aim is to develop and evaluate an explainable, context-aware, claim-level medical hallucination detection framework that integrates biomedical retrieval, natural language inference, fine-grained hallucination type classification, and clinical severity scoring."
    )
    add_numbered(
        doc,
        [
            "Develop a claim decomposition module that splits medical answers into atomic verifiable claims.",
            "Build a biomedical retrieval module that retrieves relevant evidence from PubMedQA, PubMed abstracts, or a curated biomedical corpus.",
            "Implement an NLI-based verification module that classifies each claim as supported, contradicted, or unverifiable.",
            "Design a fine-grained hallucination taxonomy for medical text and train a type classifier aligned with MedHallu categories and clinically meaningful error types.",
            "Add explainability using retrieved evidence spans, token-level highlighting, and feature attribution.",
            "Create a clinical severity score using hallucination type, contradiction strength, and clinical risk tier.",
            "Evaluate the system on MedHallu, with special attention to hard hallucination cases, and compare against binary, NLI-only, and RAG-only baselines.",
        ],
    )

    doc.add_heading("5. Dataset Plan", level=1)
    add_table(
        doc,
        ["Dataset", "Role in project"],
        [
            ["MedHallu", "Main external benchmark and target test set; includes PubMedQA-derived medical hallucination samples."],
            ["PubMedQA", "Biomedical QA and evidence source for retrieval experiments."],
            ["RAGTruth / RAGTruth++", "General RAG hallucination reference for span-level and retrieval-grounded comparison."],
            ["SciHal25", "Scientific claim verification and NLI reference."],
            ["Mu-SHROOM", "Span-level hallucination detection reference."],
            ["HealthBench", "Clinical safety and risk-rubric inspiration."],
        ],
        [2600, 6760],
    )

    doc.add_heading("6. Proposed Methodology", level=1)
    add_numbered(
        doc,
        [
            "Input medical question and generated answer.",
            "Split the answer into atomic medical claims.",
            "Retrieve top-k biomedical evidence passages for each claim.",
            "Use an NLI model to classify each claim as supported, contradicted, or unverifiable.",
            "Classify hallucination type for unsupported or contradicted claims.",
            "Generate explanations using evidence spans and token-level highlighting.",
            "Calculate clinical severity using type weight, contradiction strength, and clinical risk weight.",
        ],
    )

    doc.add_heading("7. Severity Formula", level=1)
    doc.add_paragraph("The proposed clinical severity score is:")
    p = doc.add_paragraph()
    run = p.add_run("severity_score = type_weight x contradiction_strength x clinical_risk_weight")
    run.bold = True
    doc.add_paragraph(
        "The severity output can be grouped as low, moderate, or high. High-severity examples include wrong dosage, treatment, diagnosis, contraindication, and emergency-care claims."
    )

    doc.add_heading("8. Baselines and Evaluation", level=1)
    add_table(
        doc,
        ["Baseline/System", "Purpose"],
        [
            ["Binary classifier", "Question-answer pair classified as hallucinated or not hallucinated."],
            ["NLI-only", "Claim/evidence verification without retrieval."],
            ["Retrieval-only", "Evidence similarity thresholding."],
            ["RAG + NLI", "Evidence retrieval followed by NLI verification."],
            ["Full framework", "Claim decomposition + RAG + NLI + type classification + explanation + severity."],
        ],
        [2600, 6760],
    )
    doc.add_paragraph(
        "Primary metrics will include accuracy, precision, recall, F1, macro-F1, hard-case F1 on MedHallu, per-type F1, retrieval Recall@k, and high-risk hallucination recall."
    )

    doc.add_heading("9. Work Completed So Far", level=1)
    add_bullets(
        doc,
        [
            "Collected a literature matrix with more than 30 relevant 2024-2026 papers/resources.",
            "Selected MedHallu as the main benchmark dataset.",
            "Created the research proposal, methodology, and implementation roadmap.",
            "Created an initial Python project scaffold for claim splitting, MedHallu loading, retrieval, NLI wrapping, severity scoring, dataset inspection, and binary baseline experiments.",
            "Verified an offline demo using two hand-written medical hallucination examples.",
        ],
    )

    doc.add_heading("10. Next Steps", level=1)
    add_numbered(
        doc,
        [
            "Install project requirements.",
            "Load and inspect MedHallu fields and label distributions.",
            "Run the binary baseline on pqa_labeled.",
            "Add claim-level decomposition over MedHallu answers.",
            "Build the retrieval index from available medical contexts.",
            "Run RAG + NLI verification and compare against baseline results.",
            "Add type classification and clinical severity scoring.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Appendix: Key Literature Anchors", level=1)
    add_bullets(
        doc,
        [
            "MedHallu, 2025: medical hallucination benchmark derived from PubMedQA.",
            "From RAG to Reality, 2025: coarse-grained hallucination detection via NLI fine-tuning.",
            "RT4CHART, 2026: claim decomposition and hierarchical verification for RAG hallucination.",
            "ReDeEP, 2025: mechanistic interpretability for RAG hallucination detection.",
            "RAGTruth, 2024: hallucination corpus for retrieval-augmented generation.",
            "HealthBench, 2025: healthcare evaluation benchmark with physician-created rubrics.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
